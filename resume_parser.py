import os
import re
from PyPDF2 import PdfReader
from docx import Document
from pathlib import Path


class ResumeParser:
    """Parse resume/CV files in PDF, DOCX, and TXT formats"""
    
    ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'doc'}
    
    @staticmethod
    def allowed_file(filename):
        """Check if file extension is allowed"""
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in ResumeParser.ALLOWED_EXTENSIONS
    
    @staticmethod
    def parse_pdf(filepath):
        """Extract text from PDF"""
        text = []
        try:
            with open(filepath, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    text.append(page.extract_text())
            return '\n'.join(text)
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")
    
    @staticmethod
    def parse_docx(filepath):
        """Extract text from DOCX"""
        text = []
        try:
            doc = Document(filepath)
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            return '\n'.join(text)
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")
    
    @staticmethod
    def parse_txt(filepath):
        """Extract text from TXT"""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(filepath, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            raise ValueError(f"Error parsing TXT: {str(e)}")
    
    @staticmethod
    def parse(filepath):
        """Parse any supported resume format"""
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")
        
        file_ext = Path(filepath).suffix.lower().lstrip('.')
        
        if file_ext == 'pdf':
            return ResumeParser.parse_pdf(filepath)
        elif file_ext == 'docx':
            return ResumeParser.parse_docx(filepath)
        elif file_ext in ['txt', 'doc']:
            return ResumeParser.parse_txt(filepath)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    @staticmethod
    def clean_text(text):
        """Clean and normalize resume text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep alphanumeric, spaces, and common punctuation
        text = re.sub(r'[^\w\s\-\.,/\(\)\+\#\@]', '', text)
        return text.strip()
    
    @staticmethod
    def extract_contact_info(text):
        """Extract email and phone from resume"""
        contact = {}
        
        # Email pattern
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact['email'] = email_match.group(0)
        
        # Phone pattern (various formats)
        phone_pattern = r'(?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact['phone'] = phone_match.group(0)
        
        return contact
    
    @staticmethod
    def extract_sections(text):
        """Extract common resume sections"""
        sections = {}
        
        # Common section headers
        section_keywords = {
            'summary': r'(professional\s+summary|summary|objective)',
            'experience': r'(work\s+experience|experience|employment)',
            'education': r'(education|academic)',
            'skills': r'(skills|technical\s+skills|competencies)',
            'projects': r'(projects|portfolio)',
            'certifications': r'(certifications|licenses|awards)',
            'languages': r'(languages|linguistic)'
        }
        
        # Find section starts (simple approach)
        for section_name, pattern in section_keywords.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                sections[section_name] = match.start()
        
        return sections
